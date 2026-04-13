import streamlit as st
import pandas as pd
import numpy as np
import io
import os
from zipfile import ZipFile
from docx import Document

# -----------------------------------------------------------
# PAGE CONFIG
# -----------------------------------------------------------
st.set_page_config(
    page_title="BOL Generation Tool",
    page_icon="📄",
    layout="wide"
)

# -----------------------------------------------------------
# PATHS
# -----------------------------------------------------------
STATE_CARRIER_PATH = "HD_carrier_guide.xlsx"
TEMPLATE_PATH = "BOL_template.docx"
OUTPUT_FOLDER = "BOL_created"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# -----------------------------------------------------------
# LOAD REFERENCE FILE
# -----------------------------------------------------------
@st.cache_data
def load_state_carrier():
    df = pd.read_excel(STATE_CARRIER_PATH)
    df.columns = df.columns.str.strip().str.replace(r"\s+", "", regex=True)

    tn_carrier_df = df[df["OriginState"] == "TN"].reset_index(drop=True)

    columns_to_keep = [
        "SupplierIBtoDC/StoreCarrier",
        "ResidentialDeliveryCarrier(Hd.com)",
        "DestinationState"
    ]
    tn_carrier_clean = tn_carrier_df[columns_to_keep].copy()

    mapping = {
        "AACT": "AAA Cooper Transportation",
        "EXLA": "Estes Express Lines",
        "CTII": "Central Transport Inc.",
        "ABFS": "ABF",
        "RNLO": "R&L Carriers"
    }

    tn_carrier_clean["ShippingCodeStores"] = tn_carrier_clean["SupplierIBtoDC/StoreCarrier"].map(mapping)
    tn_carrier_clean["ShippingCodeHomeDelivery"] = tn_carrier_clean["ResidentialDeliveryCarrier(Hd.com)"].map(mapping)

    return tn_carrier_clean

# -----------------------------------------------------------
# WORD TEMPLATE FILLING
# -----------------------------------------------------------
def fill_template(template_path, output_path, replacements):
    doc = Document(template_path)

    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(val))

    doc.save(output_path)

# -----------------------------------------------------------
# HELPERS
# -----------------------------------------------------------
def clean_zip(series):
    return (
        series.astype(str)
        .str.strip()
        .replace({"nan": "", "None": ""})
        .str.replace(r"\.0$", "", regex=True)
        .str.extract(r"(\d{1,5})")[0]
        .fillna("")
        .apply(lambda x: x.zfill(5) if x != "" else "")
    )

def read_chub_file(uploaded_csv):
    for skip in [4, 3, 2, 1, 0]:
        try:
            uploaded_csv.seek(0)
            df = pd.read_csv(
                uploaded_csv,
                skiprows=skip,
                encoding="utf-8",
                engine="python"
            )
            df.columns = df.columns.astype(str).str.strip().str.replace(r"\s+", "", regex=True)

            if "PONumber" in df.columns:
                return df
        except Exception:
            continue

    raise ValueError("Unable to read CommerceHub CSV correctly. Could not find expected header row.")

def merge_dn_into_ltl(df_ltl_grouped, uploaded_dn):
    df_dn = pd.read_excel(uploaded_dn)

    required_dn_cols = ["Delivery", "Sales document", "Receipt recipient"]
    missing_dn_cols = [col for col in required_dn_cols if col not in df_dn.columns]
    if missing_dn_cols:
        raise ValueError(f"DN file is missing these columns: {missing_dn_cols}")

    df_dn = df_dn[
        df_dn["Receipt recipient"].astype(str).str.contains("Home Depot", case=False, na=False)
    ].reset_index(drop=True)

    df_dn = df_dn.rename(columns={"Delivery": "DN"}).copy()

    df_dn["Sales document"] = df_dn["Sales document"].astype(str).str.strip()
    df_dn["DN"] = df_dn["DN"].astype(str).str.strip()

    df_ltl_grouped = df_ltl_grouped.copy()
    df_ltl_grouped["Sales document"] = df_ltl_grouped["Sales document"].astype(str).str.strip()

    if "DN" in df_ltl_grouped.columns:
        df_ltl_grouped = df_ltl_grouped.drop(columns=["DN"])

    merged_ltl_final = pd.merge(
        df_ltl_grouped,
        df_dn[["Sales document", "DN"]].drop_duplicates(),
        on="Sales document",
        how="left"
    )

    merged_ltl_final["DN"] = merged_ltl_final["DN"].fillna("").astype(str).str.strip()

    missing_dn_rows = merged_ltl_final[
        merged_ltl_final["DN"].eq("")
    ][["Sales document", "Purchase order no."]].drop_duplicates()

    return merged_ltl_final, missing_dn_rows

def to_zip_of_bols(df_bol):
    created_files = []

    df_bol = df_bol.reset_index(drop=True)
    progress_bar = st.progress(0.0)
    total_rows = len(df_bol)

    for idx in range(total_rows):
        row = df_bol.iloc[idx]

        dn_raw = row.get("DN", "")
        if pd.notna(dn_raw) and str(dn_raw).strip() != "":
            try:
                dn = str(int(float(dn_raw)))
            except Exception:
                dn = str(dn_raw).strip()
        else:
            dn = ""

        dn = dn.strip().replace("/", "_")

        scac = str(row.get("SCAC", "")).strip().replace("/", "_")

        weight_raw = row.get("Gross weight", "")
        if pd.notna(weight_raw) and str(weight_raw).strip() != "":
            try:
                weight = str(int(float(weight_raw)))
            except Exception:
                weight = str(weight_raw).strip()
        else:
            weight = ""

        pallet_raw = row.get("Pallet_qty", "")
        if pd.notna(pallet_raw) and str(pallet_raw).strip() != "":
            try:
                pallet_qty = str(int(float(pallet_raw)))
            except Exception:
                pallet_qty = str(pallet_raw).strip()
        else:
            pallet_qty = ""

        hd_store = row.get("HD_Store", "")
        hd_store_value = row.get("ShipToAddress1", "") if pd.notna(hd_store) and str(hd_store).strip() != "" else ""

        replacements = {
            "{{CARRIER NAME}}": row.get("Carrier_name", ""),
            "{{CUSTOMER NAME}}": row.get("ShipToName", ""),
            "{{HD_STORE}}": hd_store_value,
            "{{ADRESS}}": row.get("ShipToAddress", ""),
            "{{CITY}}": row.get("ShipToCity", ""),
            "{{STATE}}": row.get("ShipToState", ""),
            "{{ZIP CODE}}": row.get("ShipToPostalCode", ""),
            "{{PHONE NUMBER}}": row.get("ShipToDayPhone", ""),
            "{{SCAC}}": scac,
            "{{PO_NUMBER}}": row.get("PONumber", ""),
            "{{NUM_PACKAGES}}": str(row.get("Order Quantity", "")),
            "{{WEIGHT}}": weight,
            "{{CUSTOMER ORDER}}": row.get("CustomerOrderNumber", ""),
            "{{DELIVERY NUMBER}}": dn,
            "{{QTY_1}}": pallet_qty,
            "{{QTY_PACK}}": str(row.get("Order Quantity", ""))
        }

        output_file = os.path.join(OUTPUT_FOLDER, f"{dn}_{scac}.docx")
        fill_template(TEMPLATE_PATH, output_file, replacements)
        created_files.append(output_file)

        progress_bar.progress(min((idx + 1) / total_rows, 1.0))

    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, "w") as zipf:
        for file in created_files:
            zipf.write(file, os.path.basename(file))
    zip_buffer.seek(0)

    return zip_buffer, len(created_files)

# -----------------------------------------------------------
# SESSION STATE
# -----------------------------------------------------------
if "df_ltl_with_dn" not in st.session_state:
    st.session_state.df_ltl_with_dn = None

if "missing_dn_rows" not in st.session_state:
    st.session_state.missing_dn_rows = None

if "excluded_no_dn" not in st.session_state:
    st.session_state.excluded_no_dn = None

if "df_bol_preview" not in st.session_state:
    st.session_state.df_bol_preview = None

if "zip_buffer" not in st.session_state:
    st.session_state.zip_buffer = None

if "bol_count" not in st.session_state:
    st.session_state.bol_count = None
    
if "matched_chub_count" not in st.session_state:
    st.session_state.matched_chub_count = None

# -----------------------------------------------------------
# HEADER
# -----------------------------------------------------------
st.title("📄 BOL Generation Tool")
st.caption("Upload LTL, DN, and CommerceHub files to generate Bill of Lading documents.")

with st.expander("How this tool works"):
    st.markdown("""
    **Files needed**
    - **SAP LTL Cleaned Excel**
    - **DN Excel**
    - **CommerceHub CSV**

    **What the tool does**
    1. Loads the SAP LTL Cleaned file
    2. Merges the DN from the DN Excel file using **Sales document**
    3. Shows rows that did **not** get a DN
    4. Merges with CommerceHub and carrier guide
    5. Generates BOL Word files only for rows that **have a DN**
    6. Packages the BOL files into a ZIP
    """)

st.markdown("---")

# -----------------------------------------------------------
# LOAD CARRIER GUIDE
# -----------------------------------------------------------
try:
    tn_carrier_clean = load_state_carrier()
    st.success("Reference file loaded: HD_carrier_guide.xlsx")
except Exception as e:
    st.error(f"❌ Error loading carrier guide: {e}")
    st.stop()

# -----------------------------------------------------------
# FILE UPLOADS
# -----------------------------------------------------------
st.subheader("Step 1 · Upload files")

col1, col2, col3 = st.columns(3)

with col1:
    uploaded_ltl = st.file_uploader(
        "Upload SAP LTL Cleaned Excel",
        type=["xlsx"],
        key="ltl_upload"
    )

with col2:
    uploaded_dn = st.file_uploader(
        "Upload DN Excel",
        type=["xlsx", "xls"],
        key="dn_upload"
    )

with col3:
    uploaded_csv = st.file_uploader(
        "Upload CommerceHub CSV",
        type=["csv"],
        key="csv_upload"
    )

process_button = st.button("▶️ Process Files & Prepare BOL Data", use_container_width=True)

# -----------------------------------------------------------
# PROCESS
# -----------------------------------------------------------
if process_button:
    if not uploaded_ltl or not uploaded_dn or not uploaded_csv:
        st.warning("Please upload all 3 files before processing.")
    else:
        try:
            st.session_state.df_ltl_with_dn = None
            st.session_state.missing_dn_rows = None
            st.session_state.excluded_no_dn = None
            st.session_state.df_bol_preview = None
            st.session_state.zip_buffer = None
            st.session_state.bol_count = None
            st.session_state.matched_chub_count = None

            # -------------------------------
            # Load LTL and merge DN
            # -------------------------------
            df_ltl_grouped = pd.read_excel(uploaded_ltl)

            required_ltl_cols = ["Sales document", "Purchase order no."]
            missing_ltl_cols = [col for col in required_ltl_cols if col not in df_ltl_grouped.columns]
            if missing_ltl_cols:
                raise ValueError(f"LTL file is missing these columns: {missing_ltl_cols}")

            df_ltl_with_dn, missing_dn_rows = merge_dn_into_ltl(df_ltl_grouped, uploaded_dn)

            st.session_state.df_ltl_with_dn = df_ltl_with_dn
            st.session_state.missing_dn_rows = missing_dn_rows

            # -------------------------------
            # Load CHUB
            # -------------------------------
            df_csv = read_chub_file(uploaded_csv)

            required_chub_cols = [
                "PONumber",
                "ShipToName",
                "ShipToAddress1",
                "ShipToAddress2",
                "ShipToCity",
                "ShipToState",
                "ShipToPostalCode",
                "ShipToDayPhone",
                "CustomerOrderNumber"
            ]
            missing_chub_cols = [col for col in required_chub_cols if col not in df_csv.columns]
            if missing_chub_cols:
                raise ValueError(f"CommerceHub CSV is missing these columns: {missing_chub_cols}")

            df_csv["ShipToPostalCode"] = clean_zip(df_csv["ShipToPostalCode"])

            df_csv["HD_Store"] = df_csv["ShipToAddress1"].astype(str).str.extract(r"Store #(\d{3,})")

            df_csv["ShipToAddress"] = np.where(
                df_csv["ShipToAddress1"].astype(str).str.contains("THD", na=False),
                df_csv["ShipToAddress2"],
                df_csv["ShipToAddress1"]
            )

            # -------------------------------
            # Merge carriers
            # -------------------------------
            df_chub = pd.merge(
                df_csv,
                tn_carrier_clean,
                left_on="ShipToState",
                right_on="DestinationState",
                how="left"
            )

            if "DestinationState" in df_chub.columns:
                df_chub = df_chub.drop(columns=["DestinationState"])

            df_chub["SCAC"] = np.where(
                df_chub["HD_Store"].notna(),
                df_chub["SupplierIBtoDC/StoreCarrier"],
                df_chub["ResidentialDeliveryCarrier(Hd.com)"]
            )

            df_chub["Carrier_name"] = np.where(
                df_chub["HD_Store"].notna(),
                df_chub["ShippingCodeStores"],
                df_chub["ShippingCodeHomeDelivery"]
            )

            map_sap = {
                "AACT": "43564",
                "EXLA": "43558",
                "CTII": "48617",
                "ABFS": "55153",
                "RNLO": "41133"
            }

            df_chub["SAP_Carrier_Code"] = np.where(
                df_chub["HD_Store"].notna(),
                df_chub["SupplierIBtoDC/StoreCarrier"].map(map_sap),
                df_chub["ResidentialDeliveryCarrier(Hd.com)"].map(map_sap)
            )

            # -------------------------------
            # Override for CA shipments - because not present in carrier guide
            # -------------------------------
            mask_ca = df_chub["ShipToState"].astype(str).str.strip() == "CA"
            df_chub.loc[mask_ca, "SCAC"] = "EXLA"
            df_chub.loc[mask_ca, "Carrier_name"] = "Estes Express Lines"
            df_chub.loc[mask_ca, "SAP_Carrier_Code"] = "43558"

            # -------------------------------
            # Align PO types as strings
            # -------------------------------
            df_ltl_with_dn["Purchase order no."] = (
                df_ltl_with_dn["Purchase order no."]
                .astype(str)
                .str.strip()
                .str.replace(r"\.0$", "", regex=True)
            )

            df_chub["PONumber"] = (
                df_chub["PONumber"]
                .astype(str)
                .str.strip()
                .str.replace(r"\.0$", "", regex=True)
            )

            # -------------------------------
            # Count rows from LTL matched to CHUB
            # -------------------------------
            df_chub_match_check = df_chub[["PONumber"]].drop_duplicates().copy()
            df_ltl_chub_check = pd.merge(
                df_ltl_with_dn[["Purchase order no."]].copy(),
                df_chub_match_check,
                left_on="Purchase order no.",
                right_on="PONumber",
                how="left",
                indicator=True
            )
            matched_chub_count = (df_ltl_chub_check["_merge"] == "both").sum()

            ## memorixe count of pos that match from LTL cleaned and Chub
            st.session_state.matched_chub_count = matched_chub_count     

            # -------------------------------
            # Merge BOL data
            # -------------------------------
            df_bol = pd.merge(
                df_ltl_with_dn,
                df_chub,
                left_on="Purchase order no.",
                right_on="PONumber",
                how="inner"
            )

            if df_bol.empty:
                raise ValueError("No rows matched between SAP LTL file and CommerceHub CSV on PO number.")

            # Keep only rows with DN
            df_bol_with_dn = df_bol[
                df_bol["DN"].astype(str).str.strip() != ""
            ].copy()

            excluded_no_dn = df_bol[
                df_bol["DN"].astype(str).str.strip() == ""
            ][["Sales document", "Purchase order no."]].drop_duplicates()

            st.session_state.excluded_no_dn = excluded_no_dn.copy()
            st.session_state.df_bol_preview = df_bol_with_dn.copy()

            if df_bol_with_dn.empty:
                raise ValueError("No BOLs were created because none of the matched rows had a DN.")

            # -------------------------------
            # Build BOL ZIP
            # -------------------------------
            zip_buffer, bol_count = to_zip_of_bols(df_bol_with_dn)
            st.session_state.zip_buffer = zip_buffer
            st.session_state.bol_count = bol_count

            st.success("Processing completed successfully.")

        except Exception as e:
            st.error(f"❌ Error processing files: {e}")

# -----------------------------------------------------------
# RESULTS
# -----------------------------------------------------------
if st.session_state.df_ltl_with_dn is not None:
    st.markdown("---")
    st.subheader("DN merge check")

    total_rows = len(st.session_state.df_ltl_with_dn)
    missing_count = len(st.session_state.missing_dn_rows) if st.session_state.missing_dn_rows is not None else 0
    matched_count = total_rows - missing_count
    excluded_count = len(st.session_state.excluded_no_dn) if st.session_state.excluded_no_dn is not None else 0
    matched_chub_count = st.session_state.matched_chub_count if st.session_state.matched_chub_count is not None else 0
    
    m1, m2, m3, m4, m5 = st.columns(5)
    with m1:
        st.metric("Rows in LTL file", total_rows)
    with m2:
        st.metric("Rows matched to DN", matched_count)
    with m3:
        st.metric("Rows matched to CHUB", matched_chub_count)
    with m4:
        st.metric("Rows missing DN", missing_count)
    with m5:
        st.metric("Excluded from BOLs", excluded_count)

    if st.session_state.missing_dn_rows is not None and not st.session_state.missing_dn_rows.empty:
        st.warning("Some rows did not get a DN from the uploaded DN file.")
        st.dataframe(st.session_state.missing_dn_rows, use_container_width=True)
    else:
        st.success("All rows received a DN.")

if st.session_state.excluded_no_dn is not None and not st.session_state.excluded_no_dn.empty:
    st.markdown("---")
    st.subheader("Rows excluded from BOL generation")
    st.warning("These rows matched the CommerceHub file but were excluded because DN is missing.")
    st.dataframe(st.session_state.excluded_no_dn, use_container_width=True)

if st.session_state.df_bol_preview is not None:
    st.markdown("---")
    st.subheader("BOL data preview")

    preview_cols = [
        col for col in [
            "Purchase order no.",
            "Sales document",
            "DN",
            "SAP_Carrier_Code",
            "ShipToName",
            "ShipToState",
            "ShipToPostalCode",
            "SCAC",
            "Carrier_name"
        ]
        if col in st.session_state.df_bol_preview.columns
    ]

    st.dataframe(
        st.session_state.df_bol_preview[preview_cols].head(100),
        use_container_width=True
    )

if st.session_state.zip_buffer is not None:
    st.markdown("---")
    st.subheader("Download output")

    st.success(f"✅ Created {st.session_state.bol_count} BOLs.")

    st.download_button(
        "⬇️ Download All BOLs (ZIP)",
        data=st.session_state.zip_buffer,
        file_name="BOLs.zip",
        mime="application/zip",
        use_container_width=True
    )
